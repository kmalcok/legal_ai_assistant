from __future__ import annotations

import io
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TAB_CM = 6.0  # Label ":" alignment


def _clean_list_items(v: Any) -> List[str]:
    """
    Defensive cleanup for list blocks (numbered/bullets/attachments).
    - Drops empty/whitespace-only items to avoid blank numbering like "6."
    - Coerces items to strings.
    """
    items = v if isinstance(v, list) else []
    out: List[str] = []
    for it in items:
        s = str(it or "").strip()
        if not s:
            continue
        out.append(s)
    return out


def _apply_defaults(doc: Document, style: Dict[str, Any]) -> None:
    section = doc.sections[0]
    m = style.get("margins_cm", {}) if isinstance(style, dict) else {}
    section.top_margin = Cm(m.get("top", 2.5))
    section.bottom_margin = Cm(m.get("bottom", 2.5))
    section.left_margin = Cm(m.get("left", 3.75))
    section.right_margin = Cm(m.get("right", 2.5))

    normal = doc.styles["Normal"]
    font_name = style.get("font_name", "Times New Roman")
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal.font.size = Pt(style.get("font_size_pt", 12))


def _add_title(doc: Document, text: str, style: Dict[str, Any]) -> None:
    p = doc.add_paragraph(text or "")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not p.runs:
        p.add_run("")
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(style.get("heading_size_pt", 14))


def _add_header_blocks(doc: Document, blocks: List[Dict[str, Any]], style: Dict[str, Any]) -> None:
    for b in blocks:
        label = str(b.get("label") or "")
        value = str(b.get("value") or "")
        # Special case: header tags with empty label should be printed as a standalone line
        # (no ":" prefix), typically used for urgent request tags like "(... TALEPLİDİR)".
        if not label.strip():
            p = doc.add_paragraph(str(value or ""))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.line_spacing = style.get("line_spacing", 1.5)
            if not p.runs:
                p.add_run("")
            run = p.runs[0]
            run.bold = True
        else:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.line_spacing = style.get("line_spacing", 1.5)
            pf.tab_stops.add_tab_stop(Cm(TAB_CM), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.SPACES)
            pf.left_indent = Cm(TAB_CM)
            pf.first_line_indent = Cm(-TAB_CM)
            run = p.add_run(f"{label}\t: {value}")
        # IMPORTANT: Color decisions should be made by the LLM output (not by keyword heuristics),
        # otherwise unrelated parts may accidentally turn red.
        hb_style = b.get("style") if isinstance(b, dict) else None
        if isinstance(hb_style, dict) and str(hb_style.get("color") or "").strip().lower() == "red":
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _add_manual_numbered_paragraph(doc: Document, number: int, text: str, style: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = style.get("line_spacing", 1.5)
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(-0.75)
    p.add_run(f"{int(number)}. {text}")


def _add_section(doc: Document, title: str, blocks: List[Dict[str, Any]], style: Dict[str, Any]) -> None:
    t = doc.add_paragraph(title or "")
    if not t.runs:
        t.add_run("")
    tr = t.runs[0]
    tr.bold = True
    tr.font.size = Pt(style.get("font_size_pt", 12))
    # Improve readability: consistent spacing around section headings.
    try:
        t.paragraph_format.space_before = Pt(10)
        t.paragraph_format.space_after = Pt(6)
    except Exception:
        pass

    for bl in blocks:
        kind = bl.get("kind")
        text = bl.get("text")
        if kind == "paragraph":
            p = doc.add_paragraph(str(text or ""))
            p.paragraph_format.line_spacing = style.get("line_spacing", 1.5)
            p.paragraph_format.first_line_indent = Cm(1.0)
        elif kind == "block_quote":
            p = doc.add_paragraph(str(text or ""))
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.right_indent = Cm(1.0)
            p.paragraph_format.line_spacing = style.get("line_spacing", 1.5)
            try:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            except Exception:
                pass
        elif kind == "bullets":
            items = _clean_list_items(text)
            for item in items:
                p = doc.add_paragraph(str(item), style="List Bullet")
                p.paragraph_format.line_spacing = style.get("line_spacing", 1.5)
        elif kind == "numbered":
            items = _clean_list_items(text)
            for i, item in enumerate(items, start=1):
                _add_manual_numbered_paragraph(doc, i, str(item), style)
        else:
            p = doc.add_paragraph(str(text or ""))
            p.paragraph_format.line_spacing = style.get("line_spacing", 1.5)


def _add_signature(doc: Document, sig: Dict[str, Any], style: Dict[str, Any]) -> None:
    if isinstance(sig, dict) and sig.get("place_date_line"):
        p = doc.add_paragraph(str(sig.get("place_date_line") or ""))
        p.paragraph_format.line_spacing = style.get("line_spacing", 1.5)
    phrase = str((sig or {}).get("phrase") or "").strip()
    name_line = str((sig or {}).get("name_line") or "").strip()
    if phrase:
        p2 = doc.add_paragraph(phrase)
        p2.paragraph_format.line_spacing = style.get("line_spacing", 1.5)
    if name_line:
        p3 = doc.add_paragraph(name_line)
        p3.paragraph_format.line_spacing = style.get("line_spacing", 1.5)


def _add_attachments(doc: Document, attachments: List[str], style: Dict[str, Any]) -> None:
    attachments_clean = _clean_list_items(attachments)
    if not attachments_clean:
        return
    doc.add_paragraph("")
    t = doc.add_paragraph("EKLER")
    if not t.runs:
        t.add_run("")
    t.runs[0].bold = True
    for i, a in enumerate(attachments_clean, start=1):
        _add_manual_numbered_paragraph(doc, i, str(a), style)


def render_docx_bytes(output: Dict[str, Any]) -> bytes:
    """
    Render DOCX bytes from an output JSON matching petitions/output.schema.json.
    """
    meta = output.get("meta") if isinstance(output.get("meta"), dict) else {}
    style = meta.get("style") if isinstance(meta.get("style"), dict) else {}

    doc = Document()
    _apply_defaults(doc, style)

    _add_title(doc, str(meta.get("court") or ""), style)
    # Place urgent header tags immediately under the main title.
    # Preferred source: `urgent_tags` (explicit, schema field).
    # Back-compat: also accept header_blocks with empty label.
    urgent_tags = output.get("urgent_tags") or []
    urgent_values: List[str] = []
    if isinstance(urgent_tags, list):
        for v in urgent_tags:
            s = str(v or "").strip()
            if s:
                urgent_values.append(s)

    blocks = output.get("header_blocks") or []
    if not isinstance(blocks, list):
        blocks = []
    urgent_blocks: List[Dict[str, Any]] = []
    rest_blocks: List[Dict[str, Any]] = []
    for b in blocks:
        if isinstance(b, dict) and not str(b.get("label") or "").strip() and str(b.get("value") or "").strip():
            urgent_blocks.append(b)
        else:
            rest_blocks.append(b)

    # Merge: urgent_tags (preferred) + back-compat urgent header blocks.
    if urgent_values:
        seen = set()
        merged = []
        for s in urgent_values:
            if s in seen:
                continue
            seen.add(s)
            merged.append({"label": "", "value": s})
        urgent_blocks = merged + urgent_blocks

    if urgent_blocks:
        doc.add_paragraph("")
        _add_header_blocks(doc, urgent_blocks, style)
        doc.add_paragraph("")

    # Header blocks (party/date/file info etc.)
    _add_header_blocks(doc, rest_blocks, style)
    doc.add_paragraph("")

    for s in output.get("sections") or []:
        if not isinstance(s, dict):
            continue
        _add_section(doc, str(s.get("title") or ""), s.get("blocks") or [], style)
        doc.add_paragraph("")

    _add_signature(doc, output.get("signature") or {}, style)
    _add_attachments(doc, output.get("attachments") or [], style)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


