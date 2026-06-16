from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
from defusedxml import ElementTree as DefusedET
from pypdf import PdfReader

import docx

from .utils import normalize_text, split_into_pseudopages


@dataclass
class Page:
    page_no: int
    text: str


@dataclass
class ExtractedDocument:
    file_type: str  # pdf|docx|udf|unknown
    pages: List[Page]
    meta: Dict[str, Any]


class DocumentExtractError(Exception):
    pass


def _is_zip_bytes(b: bytes) -> bool:
    return len(b) >= 4 and b[:4] == b"PK\x03\x04"


def _find_zip_suffix(names: list[str], suffix: str) -> Optional[str]:
    suffix = suffix.lower()
    for n in names:
        if n.lower().endswith(suffix):
            return n
    return None


def _collect_xml_text(xml_bytes: bytes) -> str:
    root = DefusedET.fromstring(xml_bytes)
    parts: list[str] = []
    for elem in root.iter():
        if elem.text and elem.text.strip():
            parts.append(elem.text.strip())
        if elem.tail and elem.tail.strip():
            parts.append(elem.tail.strip())
    return normalize_text("\n".join(parts))


def extract_pdf_pages(file_bytes: bytes) -> Tuple[List[Page], Dict[str, Any]]:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[Page] = []
    used_fitz_fallback = False
    for i, p in enumerate(reader.pages, start=1):
        try:
            t = p.extract_text() or ""
        except Exception:
            t = ""
        pages.append(Page(page_no=i, text=normalize_text(t)))

    if pages and not any(page.text for page in pages):
        try:
            fitz_doc = fitz.open(stream=file_bytes, filetype="pdf")
            try:
                fallback_pages: list[Page] = []
                for i, page in enumerate(fitz_doc, start=1):
                    fallback_pages.append(Page(page_no=i, text=normalize_text(page.get_text("text") or "")))
                if any(page.text for page in fallback_pages):
                    pages = fallback_pages
                    used_fitz_fallback = True
            finally:
                fitz_doc.close()
        except Exception:
            pass

    if pages and not any(page.text for page in pages):
        raise DocumentExtractError("PDF text could not be extracted. The file may be image-only or require OCR.")

    meta: Dict[str, Any] = {"pages": len(reader.pages), "used_fitz_fallback": bool(used_fitz_fallback)}
    return pages, meta


def extract_docx_pages(file_bytes: bytes, *, pseudo_page_chars: int = 5000) -> Tuple[List[Page], Dict[str, Any]]:
    d = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for p in d.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())

    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full = normalize_text("\n".join(parts))
    chunks = split_into_pseudopages(full, max_chars=pseudo_page_chars)
    pages = [Page(page_no=i + 1, text=normalize_text(ch)) for i, ch in enumerate(chunks)]
    return pages, {"pseudo_pages": True, "pseudo_page_chars": int(pseudo_page_chars)}


def extract_udf_pages(file_bytes: bytes, *, pseudo_page_chars: int = 5000) -> Tuple[List[Page], Dict[str, Any]]:
    if not _is_zip_bytes(file_bytes):
        raise DocumentExtractError("UDF expected ZIP-like bytes (PK header not found).")
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            names = zf.namelist()
            content_name = _find_zip_suffix(names, "content.xml")
            props_name = _find_zip_suffix(names, "documentproperties.xml")
            sign_name = _find_zip_suffix(names, "sign.sgn")
            if not content_name:
                raise DocumentExtractError("UDF content.xml not found.")
            content_xml = zf.read(content_name)
            props_xml = zf.read(props_name) if props_name else None
    except zipfile.BadZipFile as exc:
        raise DocumentExtractError(f"Invalid UDF zip: {exc}") from exc

    text = _collect_xml_text(content_xml)
    chunks = split_into_pseudopages(text, max_chars=pseudo_page_chars)
    pages = [Page(page_no=i + 1, text=normalize_text(ch)) for i, ch in enumerate(chunks)]

    properties: Dict[str, Any] = {}
    if props_xml:
        try:
            root = DefusedET.fromstring(props_xml)
            for e in root.iter():
                if e is root:
                    continue
                if e.text and e.text.strip():
                    properties[e.tag] = e.text.strip()
        except Exception:
            properties = {}

    meta: Dict[str, Any] = {
        "has_signature": bool(sign_name),
        "properties": properties,
        "pseudo_pages": True,
        "pseudo_page_chars": int(pseudo_page_chars),
    }
    return pages, meta


def extract_document(
    file_bytes: bytes,
    filename: str,
    *,
    pseudo_page_chars: int = 5000,
) -> ExtractedDocument:
    ext = (filename or "").lower().rsplit(".", 1)[-1] if "." in (filename or "") else ""
    ext = ext.strip().lower()

    if ext == "pdf":
        pages, meta = extract_pdf_pages(file_bytes)
        return ExtractedDocument(file_type="pdf", pages=pages, meta=meta)
    if ext == "docx":
        pages, meta = extract_docx_pages(file_bytes, pseudo_page_chars=pseudo_page_chars)
        return ExtractedDocument(file_type="docx", pages=pages, meta=meta)
    if ext == "udf":
        pages, meta = extract_udf_pages(file_bytes, pseudo_page_chars=pseudo_page_chars)
        return ExtractedDocument(file_type="udf", pages=pages, meta=meta)

    # sniff zip-like docx / udf
    if _is_zip_bytes(file_bytes):
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                names = set(n.lower() for n in zf.namelist())
            if "[content_types].xml" in names:
                pages, meta = extract_docx_pages(file_bytes, pseudo_page_chars=pseudo_page_chars)
                return ExtractedDocument(file_type="docx", pages=pages, meta=meta)
            if any(n.endswith("content.xml") for n in names):
                pages, meta = extract_udf_pages(file_bytes, pseudo_page_chars=pseudo_page_chars)
                return ExtractedDocument(file_type="udf", pages=pages, meta=meta)
        except Exception:
            pass

    raise DocumentExtractError("Unsupported document type (expected pdf/docx/udf).")


