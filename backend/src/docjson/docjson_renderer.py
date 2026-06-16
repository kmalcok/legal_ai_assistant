from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


@dataclass(frozen=True)
class RenderedDocx:
    filename: str
    mime: str
    content: bytes


def _now_utc_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _get(d: Dict[str, Any], key: str, default: Any = None) -> Any:
    v = d.get(key, default)
    return default if v is None else v


def _ensure_nonempty(s: str, *, fallback: str) -> str:
    s = (s or "").strip()
    return s if s else fallback


def _apply_defaults(doc: Document, spec: Dict[str, Any]) -> None:
    defaults = spec.get("defaults") if isinstance(spec.get("defaults"), dict) else {}
    font = defaults.get("font") if isinstance(defaults.get("font"), dict) else {}
    para = defaults.get("paragraph") if isinstance(defaults.get("paragraph"), dict) else {}

    normal = doc.styles["Normal"]
    font_name = str(font.get("name") or "Calibri")
    font_size = float(font.get("sizePt") or 11)
    normal.font.name = font_name
    # Ensure consistent font mapping for non-ASCII scripts too.
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        normal._element.rPr.rFonts.set(qn("w:cs"), font_name)
    except Exception:
        pass
    normal.font.size = Pt(font_size)

    # Keep heading fonts consistent with Normal to avoid viewer/font fallbacks
    # that can make Turkish characters look "broken".
    for heading_style in ("Heading 1", "Heading 2", "Title"):
        try:
            st = doc.styles[heading_style]
            st.font.name = font_name
            st._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            st._element.rPr.rFonts.set(qn("w:cs"), font_name)
            st.font.size = Pt(font_size if heading_style != "Heading 1" else max(font_size + 2, font_size))
        except Exception:
            pass

    # store paragraph defaults for later use
    doc._docjson_defaults = {  # type: ignore[attr-defined]
        "line_spacing": float(para.get("lineSpacing") or 1.15),
        "space_after_pt": float(para.get("spaceAfterPt") or 6),
    }


def _para_defaults(doc: Document) -> Dict[str, float]:
    d = getattr(doc, "_docjson_defaults", None)
    if isinstance(d, dict):
        return {"line_spacing": float(d.get("line_spacing", 1.15)), "space_after_pt": float(d.get("space_after_pt", 6))}
    return {"line_spacing": 1.15, "space_after_pt": 6.0}


def _set_alignment(p, alignment: str | None) -> None:
    a = (alignment or "").strip().lower()
    if a == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif a == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_heading(doc: Document, block: Dict[str, Any]) -> None:
    lvl = int(block.get("level") or 1)
    txt = str(block.get("text") or "").strip()
    if not txt:
        return
    style = "Heading 1" if lvl == 1 else "Heading 2"
    p = doc.add_paragraph(txt, style=style)
    d = _para_defaults(doc)
    p.paragraph_format.line_spacing = d["line_spacing"]
    p.paragraph_format.space_after = Pt(d["space_after_pt"])


def _add_paragraph(doc: Document, block: Dict[str, Any]) -> None:
    runs = block.get("runs")
    if not isinstance(runs, list) or not runs:
        return  # hard rule: empty paragraphs not allowed
    p = doc.add_paragraph()
    d = _para_defaults(doc)
    p.paragraph_format.line_spacing = d["line_spacing"]
    p.paragraph_format.space_after = Pt(d["space_after_pt"])
    _set_alignment(p, block.get("alignment"))
    for r in runs:
        if not isinstance(r, dict):
            continue
        t = str(r.get("text") or "")
        if t == "":
            continue
        run = p.add_run(t)
        run.bold = bool(r.get("bold") or False)
        run.italic = bool(r.get("italic") or False)
        run.underline = bool(r.get("underline") or False)


def _add_list(doc: Document, block: Dict[str, Any], *, level: int = 0) -> None:
    ordered = bool(block.get("ordered") or False)
    items = block.get("items")
    if not isinstance(items, list) or not items:
        return
    style = "List Number" if ordered else "List Bullet"
    d = _para_defaults(doc)

    def add_item(item: Dict[str, Any], lvl: int) -> None:
        text = str(item.get("text") or "").strip()
        if text:
            p = doc.add_paragraph(text, style=style)
            p.paragraph_format.line_spacing = d["line_spacing"]
            p.paragraph_format.space_after = Pt(d["space_after_pt"])
            # crude nesting indent
            if lvl > 0:
                p.paragraph_format.left_indent = Pt(18 * lvl)
        children = item.get("children")
        if isinstance(children, list):
            for ch in children:
                if isinstance(ch, dict):
                    add_item(ch, lvl + 1)

    for it in items:
        if isinstance(it, dict):
            add_item(it, level)


def _add_table(doc: Document, block: Dict[str, Any]) -> None:
    rows = block.get("rows")
    if not isinstance(rows, list):
        return
    header = block.get("header")
    ncols = None
    if isinstance(header, list) and header:
        ncols = len(header)
    elif rows:
        first = rows[0]
        if isinstance(first, list):
            ncols = len(first)
    if not ncols or ncols <= 0:
        return

    tbl = doc.add_table(rows=0, cols=int(ncols))
    tbl.style = str(block.get("style") or "Table Grid")

    def add_row(vals: List[Any], is_header: bool = False) -> None:
        tr = tbl.add_row().cells
        for i in range(int(ncols)):
            v = ""
            if i < len(vals):
                v = str(vals[i] if vals[i] is not None else "")
            tr[i].text = v
            if is_header and tr[i].paragraphs and tr[i].paragraphs[0].runs:
                tr[i].paragraphs[0].runs[0].bold = True

    if isinstance(header, list) and header:
        add_row(header, is_header=True)

    for r in rows:
        if isinstance(r, list):
            add_row(r, is_header=False)


def _add_hr(doc: Document) -> None:
    # Word doesn't have a simple "hr"; we emulate with a divider line.
    p = doc.add_paragraph("―" * 24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_provenance(doc: Document, block: Dict[str, Any]) -> None:
    items = block.get("items")
    if not isinstance(items, list) or not items:
        return
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for it in items:
        if not isinstance(it, dict):
            continue
        lab = str(it.get("label") or "").strip()
        val = str(it.get("value") or "").strip()
        if not lab and not val:
            continue
        row = tbl.add_row().cells
        row[0].text = lab
        row[1].text = val
        if row[0].paragraphs and row[0].paragraphs[0].runs:
            row[0].paragraphs[0].runs[0].bold = True


def _render_blocks(doc: Document, blocks: List[Dict[str, Any]]) -> None:
    for b in blocks:
        if not isinstance(b, dict):
            continue
        t = str(b.get("type") or "").strip()
        if t == "heading":
            _add_heading(doc, b)
        elif t == "paragraph":
            _add_paragraph(doc, b)
        elif t == "list":
            _add_list(doc, b)
        elif t == "table":
            _add_table(doc, b)
        elif t == "pageBreak":
            doc.add_page_break()
        elif t == "hr":
            _add_hr(doc)
        elif t == "provenance":
            _add_provenance(doc, b)
        else:
            # ignore unknown block types (hard rule: do not hallucinate)
            continue


def _apply_events(events: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Apply append-only events to build a final DOC-JSON spec.
    """
    spec: Dict[str, Any] = {}
    blocks: List[Dict[str, Any]] = []
    inited = False
    for ev in events:
        if not isinstance(ev, dict):
            continue
        op = str(ev.get("op") or "").strip().lower()
        if op == "init":
            meta = ev.get("meta") if isinstance(ev.get("meta"), dict) else {}
            title = str(meta.get("title") or "").strip()
            spec = {
                "schemaVersion": "1.0",
                "meta": {
                    "docId": "",
                    "title": title,
                    "language": "tr-TR",
                    "createdAt": _now_utc_iso(),
                    "generator": {"agent": "main-agent", "version": "1.0"},
                },
                "defaults": {"font": {"name": "Calibri", "sizePt": 11}, "paragraph": {"lineSpacing": 1.15, "spaceAfterPt": 6}},
                "styles": {"Title": {}, "H1": {}, "H2": {}, "Normal": {}},
                "body": {"blocks": blocks},
            }
            inited = True
        elif op == "append":
            if not inited:
                # ignore appends before init
                continue
            path = str(ev.get("path") or "")
            if path != "body.blocks":
                continue
            val = ev.get("value")
            if isinstance(val, dict):
                blocks.append(val)
    return spec


def render_docjson_to_docx(spec_or_events: Any, *, filename: Optional[str] = None) -> RenderedDocx:
    """
    Accept either:
    - full DOC-JSON object (schemaVersion/meta/body.blocks)
    - list of append-only events (init + append)
    """
    spec: Dict[str, Any]
    if isinstance(spec_or_events, list):
        spec = _apply_events([x for x in spec_or_events if isinstance(x, dict)])
    elif isinstance(spec_or_events, dict):
        # if it's a single event object, wrap into a list
        if "op" in spec_or_events:
            spec = _apply_events([spec_or_events])  # likely incomplete but deterministic
        else:
            spec = spec_or_events
    else:
        raise ValueError("invalid_docjson: must be object or event list")

    meta = spec.get("meta") if isinstance(spec.get("meta"), dict) else {}
    title = str(meta.get("title") or "").strip()
    out_name = filename or (title.strip() + ".docx" if title else "document.docx")
    out_name = _ensure_nonempty(out_name, fallback="document.docx")
    if not out_name.lower().endswith(".docx"):
        out_name = out_name + ".docx"

    body = spec.get("body") if isinstance(spec.get("body"), dict) else {}
    blocks = body.get("blocks") if isinstance(body.get("blocks"), list) else []

    doc = Document()
    _apply_defaults(doc, spec)
    _render_blocks(doc, [b for b in blocks if isinstance(b, dict)])

    buf = io.BytesIO()
    doc.save(buf)
    return RenderedDocx(filename=out_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", content=buf.getvalue())


